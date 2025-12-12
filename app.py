"""
Streamlined Radiology Reporting Assistant
Focus: Templates, Technique, and Structured Reporting
WITH Perplexity AI Integration
"""

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
import json
import os
import datetime
import hashlib
from collections import defaultdict
from perplexityai import PerplexitySync
from dotenv import load_dotenv

# ===== LOAD ENVIRONMENT VARIABLES =====
load_dotenv()

# ===== CONFIGURATION =====
TEMPLATES_FILE = "saved_templates.json"
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")

if not PERPLEXITY_API_KEY:
    st.warning("⚠️ PERPLEXITY_API_KEY not found in .env file")

# ===== PERPLEXITY AI INTEGRATION =====
class PerplexityAIHelper:
    """Helper class for Perplexity AI integration."""
    
    def __init__(self, api_key=None):
        self.api_key = api_key or PERPLEXITY_API_KEY
        self.client = None
        if self.api_key:
            try:
                self.client = PerplexitySync(api_key=self.api_key)
                print("✅ Perplexity AI Connected!")
            except Exception as e:
                print(f"❌ Failed to initialize Perplexity AI: {str(e)}")
    
    def is_available(self):
        """Check if Perplexity AI is available."""
        return self.client is not None
    
    def generate_report_from_findings(self, findings_text, modality="MRI", contrast="Without contrast", style="Standard"):
        """Generate a complete radiology report from findings."""
        if not self.client or not findings_text.strip():
            return None
        
        try:
            # Define style prompts
            style_prompts = {
                "Standard": "Create a standard professional radiology report.",
                "Detailed": "Provide a detailed report with comprehensive descriptions.",
                "Concise": "Be very concise while covering essential findings.",
                "Teaching": "Include teaching points and explanations suitable for trainees."
            }
            
            prompt = f"""You are a senior board-certified radiologist. Generate a structured radiology report based on these findings:

MODALITY: {modality}
CONTRAST: {contrast}

CLINICAL FINDINGS PROVIDED:
{findings_text}

{style_prompts.get(style, "Create a standard professional radiology report.")}

Please generate a complete radiology report with the following structure:

TECHNIQUE:
- Briefly describe the imaging technique

FINDINGS:
- Organize findings by anatomical region/system
- Use precise radiological terminology
- Include measurements when appropriate
- Note any comparisons with prior studies

IMPRESSION:
- Numbered conclusions (1-4 points)
- Clinical recommendations
- Follow-up suggestions if indicated

Use proper medical terminology. Format with clear headings. Be professional and accurate."""
            
            response = self.client.query(
                model="sonar",
                messages=[
                    {"role": "system", "content": "You are a senior radiologist assistant specializing in creating structured, professional radiology reports. Always use proper medical terminology and follow standard report formatting."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000
            )
            
            if response.choices:
                return response.choices[0].message.content
            return None
            
        except Exception as e:
            st.error(f"AI generation error: {str(e)}")
            return None
    
    def generate_impression_only(self, findings_text):
        """Generate only the IMPRESSION section from findings."""
        if not self.client:
            return None
        
        try:
            prompt = f"""Based on these radiology findings:
{findings_text}

Generate a concise IMPRESSION section with:
1. Numbered conclusions (2-4 points)
2. Clinical recommendations
3. Follow-up suggestions if needed

Use bullet points or numbered list. Be clinically actionable."""
            
            response = self.client.query(
                model="sonar",
                messages=[
                    {"role": "system", "content": "You are a radiologist creating concise, clinically useful impression sections."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500
            )
            
            if response.choices:
                return response.choices[0].message.content
            return None
            
        except Exception as e:
            st.error(f"AI impression generation error: {str(e)}")
            return None
    
    def enhance_existing_report(self, report_text):
        """Enhance an existing report with better terminology."""
        if not self.client:
            return None
        
        try:
            prompt = f"""Improve this radiology report with better medical terminology and structure while preserving all findings:

{report_text}

Make it more professional and structured but don't add or remove any findings. Improve formatting and organization."""
            
            response = self.client.query(
                model="sonar-pro",
                messages=[
                    {"role": "system", "content": "You are a senior radiologist editing reports for better medical terminology and structure."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2500
            )
            
            if response.choices:
                return response.choices[0].message.content
            return None
            
        except Exception as e:
            st.error(f"AI enhancement error: {str(e)}")
            return None
    
    def generate_differential_diagnosis(self, findings_text):
        """Generate differential diagnosis from findings."""
        if not self.client:
            return None
        
        try:
            prompt = f"""Based on these radiology findings:
{findings_text}

Provide a structured differential diagnosis with:
1. MOST LIKELY DIAGNOSIS (with confidence level and key features)
2. ALTERNATIVE DIAGNOSES (list 2-3 possibilities)
3. KEY DISTINGUISHING FEATURES
4. RECOMMENDED NEXT STEPS (imaging/lab/clinical)"""
            
            response = self.client.query(
                model="sonar-reasoning",
                messages=[
                    {"role": "system", "content": "You are a radiologist providing structured differential diagnoses."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000
            )
            
            if response.choices:
                return response.choices[0].message.content
            return None
            
        except Exception as e:
            st.error(f"AI differential diagnosis error: {str(e)}")
            return None
    
    def generate_findings_suggestions(self, brief_description):
        """Generate detailed findings from a brief description."""
        if not self.client:
            return None
        
        try:
            prompt = f"""As a radiologist, expand this brief description into detailed, structured radiology findings:

BRIEF: {brief_description}

Provide:
1. Systematic findings organized by anatomy
2. Specific measurements if applicable
3. Comparisons if previous studies mentioned
4. Proper radiology terminology"""
            
            response = self.client.query(
                model="sonar",
                messages=[
                    {"role": "system", "content": "You are a radiologist expanding brief descriptions into detailed findings."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000
            )
            
            if response.choices:
                return response.choices[0].message.content
            return None
            
        except Exception as e:
            st.error(f"AI findings generation error: {str(e)}")
            return None

# ===== TEMPLATE SYSTEM =====
class TemplateSystem:
    """Template system for managing report templates."""
    
    def __init__(self):
        self.templates = {}
        self.load_templates()
    
    def load_templates(self):
        """Load templates from file."""
        if os.path.exists(TEMPLATES_FILE):
            try:
                with open(TEMPLATES_FILE, 'r') as f:
                    self.templates = json.load(f)
            except:
                self.templates = {}
    
    def save_templates(self):
        """Save templates to file."""
        try:
            with open(TEMPLATES_FILE, 'w') as f:
                json.dump(self.templates, f, indent=2)
        except:
            pass
    
    def add_template(self, name, content, template_type="findings"):
        """Add a new template."""
        self.templates[name] = {
            "content": content,
            "type": template_type,
            "created_by": st.session_state.get('current_user', 'unknown'),
            "created_date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
            "used_count": 0
        }
        self.save_templates()
    
    def get_template_heading(self, template_name):
        """Get heading format for a template."""
        if template_name not in self.templates:
            return template_name.upper()
        
        template = self.templates[template_name]
        template_type = template.get("type", "findings")
        
        heading_map = {
            "technique": "TECHNIQUE",
            "findings": "FINDINGS",
            "impression": "IMPRESSION"
        }
        
        return heading_map.get(template_type, template_name.upper())
    
    def apply_template(self, template_name, current_text=""):
        """Apply a template with proper heading."""
        if template_name not in self.templates:
            return current_text
        
        template = self.templates[template_name]
        heading = self.get_template_heading(template_name)
        
        # Increment usage count
        template["used_count"] = template.get("used_count", 0) + 1
        self.save_templates()
        
        # Format with heading
        formatted_template = f"\n{heading}:\n{template['content']}"
        
        if current_text:
            return current_text + formatted_template
        return formatted_template
    
    def get_user_templates(self, username):
        """Get templates created by specific user."""
        user_templates = {}
        for name, data in self.templates.items():
            if data.get("created_by") == username:
                user_templates[name] = data
        return user_templates

def create_word_document(report_text, technique_info, report_date):
    """Create a Word document with proper formatting."""
    doc = Document()
    
    # Title
    title = doc.add_heading('RADIOLOGY REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Add technique section
    doc.add_heading('TECHNIQUE', level=1)
    if technique_info:
        doc.add_paragraph(f"Modality: {technique_info.get('modality', 'Not specified')}")
        doc.add_paragraph(f"Contrast: {technique_info.get('contrast', 'Without contrast')}")
        if technique_info.get('sequences'):
            doc.add_paragraph(f"Protocol: {technique_info['sequences']}")
    
    # Add spacing
    doc.add_paragraph()
    
    # Parse and add report sections
    lines = report_text.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Check if line is a heading (uppercase and ends with colon)
        if line_stripped.endswith(':') and line_stripped[:-1].replace(' ', '').isupper():
            heading_text = line_stripped[:-1]
            doc.add_heading(heading_text, level=1)
        else:
            # Regular content
            doc.add_paragraph(line_stripped)
    
    # Add footer
    doc.add_page_break()
    doc.add_heading('REPORT DETAILS', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Generated by: {st.session_state.get('current_user', 'Unknown')}\n")
    p.add_run(f"Generation date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    # Add AI notice if applicable
    if st.session_state.get('ai_generated', False):
        p.add_run("AI-assisted using Perplexity AI")
    
    return doc

def ensure_proper_structure(text):
    """Ensure report has proper FINDINGS and IMPRESSION structure."""
    if not text:
        return text
    
    # Ensure FINDINGS section exists
    if "FINDINGS:" not in text.upper():
        lines = text.split('\n')
        # Find where to insert FINDINGS
        insert_index = 0
        for i, line in enumerate(lines):
            if line.strip().startswith("TECHNIQUE:"):
                insert_index = i + 1
                while insert_index < len(lines) and lines[insert_index].strip():
                    insert_index += 1
                break
        
        if insert_index < len(lines):
            lines.insert(insert_index, "\nFINDINGS:\n")
        else:
            lines.append("\nFINDINGS:\n")
        
        text = '\n'.join(lines)
    
    # Ensure IMPRESSION section exists
    if "IMPRESSION:" not in text.upper():
        text += "\n\nIMPRESSION:\nFindings as described above. Clinical correlation recommended."
    
    return text

def parse_findings_for_ai(finding_text):
    """Extract just the findings section for AI processing."""
    if not finding_text:
        return finding_text
    
    # Try to find FINDINGS section
    lines = finding_text.split('\n')
    findings_start = -1
    findings_end = -1
    
    for i, line in enumerate(lines):
        if "FINDINGS:" in line.upper():
            findings_start = i
            break
    
    if findings_start >= 0:
        # Look for next section heading
        for i in range(findings_start + 1, len(lines)):
            if any(section in lines[i].upper() for section in ["IMPRESSION:", "TECHNIQUE:", "RECOMMENDATION:"]):
                findings_end = i
                break
        
        if findings_end == -1:
            findings_end = len(lines)
        
        # Extract findings content
        findings_content = '\n'.join(lines[findings_start+1:findings_end])
        return findings_content.strip()
    
    # If no FINDINGS section found, return all text
    return finding_text

def clean_ai_response(ai_text):
    """Clean and format AI response."""
    if not ai_text:
        return ai_text
    
    # Remove any markdown formatting
    ai_text = ai_text.replace('**', '').replace('*', '').replace('`', '')
    
    # Ensure proper section headings
    ai_text = ai_text.replace('Technique:', 'TECHNIQUE:')
    ai_text = ai_text.replace('Findings:', 'FINDINGS:')
    ai_text = ai_text.replace('Impression:', 'IMPRESSION:')
    
    # Add TECHNIQUE if missing
    if "TECHNIQUE:" not in ai_text.upper():
        tech_info = st.session_state.technique_info
        technique_section = f"TECHNIQUE:\n"
        technique_section += f"Modality: {tech_info.get('modality', 'Not specified')}\n"
        technique_section += f"Contrast: {tech_info.get('contrast', 'Without contrast')}\n"
        if tech_info.get('sequences'):
            technique_section += f"Protocol: {tech_info['sequences']}\n"
        
        # Insert at beginning
        ai_text = technique_section + "\n" + ai_text
    
    return ai_text

# ===== INITIALIZE SESSION STATE =====
def init_session_state():
    """Initialize session state variables."""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.users = {
            "admin": {"password": hashlib.sha256("admin123".encode()).hexdigest(), "role": "admin"},
            "radiologist": {"password": hashlib.sha256("rad123".encode()).hexdigest(), "role": "radiologist"}
        }
        st.session_state.report_history = []
        st.session_state.report_draft = ""
        st.session_state.ai_report = ""
        st.session_state.report_date = ""
        st.session_state.current_user = "default"
        st.session_state.logged_in = False
        st.session_state.template_system = TemplateSystem()
        st.session_state.technique_info = {
            "modality": "MRI",
            "contrast": "Without contrast",
            "sequences": "Standard sequences"
        }
        st.session_state.ai_helper = PerplexityAIHelper(PERPLEXITY_API_KEY)
        st.session_state.ai_generated = False
        st.session_state.ai_style = "Standard"
        st.session_state.ai_mode = "Generate Full Report"

# ===== STREAMLIT APP =====
def main():
    # Page config
    st.set_page_config(
        page_title="Radiology Reporting Assistant",
        layout="wide",
        page_icon="🏥"
    )
    
    # Initialize session state
    init_session_state()
    
    # ===== LOGIN SYSTEM =====
    if not st.session_state.logged_in:
        st.title("🔐 Radiology Assistant - Login")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            with st.container(border=True):
                st.subheader("Login")
                username = st.text_input("Username", key="login_user")
                password = st.text_input("Password", type="password", key="login_pass")
                
                if st.button("Login", type="primary", use_container_width=True):
                    hashed_pw = hashlib.sha256(password.encode()).hexdigest()
                    if username in st.session_state.users:
                        if st.session_state.users[username]["password"] == hashed_pw:
                            st.session_state.logged_in = True
                            st.session_state.current_user = username
                            
                            # Check AI availability
                            if st.session_state.ai_helper and st.session_state.ai_helper.is_available():
                                st.success(f"✅ Welcome, {username}! AI Assistant is ready.")
                            else:
                                st.warning(f"⚠️ Welcome, {username}! AI features are unavailable. Check your API key.")
                            
                            st.rerun()
                        else:
                            st.error("Invalid password")
                    else:
                        st.error("User not found")
                
                st.info("Demo: admin/admin123 or radiologist/rad123")
        
        return
    
    # ===== MAIN APPLICATION =====
    
    # Header
    st.title("🏥 Radiology Reporting Assistant 🤖")
    
    # AI Status Banner
    if st.session_state.ai_helper and st.session_state.ai_helper.is_available():
        st.success("✅ AI Assistant is active and ready to generate reports!")
    else:
        st.error("⚠️ AI Assistant is unavailable. Please check your .env file and ensure PERPLEXITY_API_KEY is set correctly.")
    
    col_user1, col_user2 = st.columns([3, 1])
    with col_user1:
        st.markdown(f"**User:** {st.session_state.current_user}")
    
    with col_user2:
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.rerun()
    
    # Main columns
    col1, col2 = st.columns([1, 1])
    
    with col1:
        # ===== LEFT PANEL: INPUT & TEMPLATES =====
        st.header("✍️ Report Creation")
        
        # Technique Information
        with st.expander("🔬 Technique Details", expanded=True):
            col_tech1, col_tech2 = st.columns(2)
            with col_tech1:
                modality = st.selectbox(
                    "Modality",
                    ["MRI", "CT", "Ultrasound", "X-ray", "PET-CT", "Mammography"],
                    key="modality_select"
                )
                
                contrast = st.selectbox(
                    "Contrast Administration",
                    ["Without contrast", "With contrast", "With and without contrast", "Not specified"],
                    key="contrast_select"
                )
            
            with col_tech2:
                sequences = st.text_area(
                    "Sequences/Protocol",
                    value=st.session_state.technique_info.get('sequences', ''),
                    placeholder="e.g., T1, T2, FLAIR, DWI, ADC",
                    key="sequences_input",
                    height=80
                )
            
            if st.button("💾 Save Technique", key="save_tech_button", use_container_width=True):
                st.session_state.technique_info = {
                    "modality": modality,
                    "contrast": contrast,
                    "sequences": sequences if sequences else "Standard sequences"
                }
                st.success("Technique details saved!")
        
        # AI Assistance Panel (only show if AI is available)
        if st.session_state.ai_helper and st.session_state.ai_helper.is_available():
            st.divider()
            st.header("🤖 AI Assistance")
            
            ai_mode = st.radio(
                "AI Mode:",
                ["Generate Full Report", "Enhance Existing", "Differential Diagnosis", "Generate Findings Only"],
                horizontal=True,
                key="ai_mode_radio"
            )
            st.session_state.ai_mode = ai_mode
            
            if ai_mode == "Generate Full Report":
                st.caption("AI will generate a complete structured report from your findings")
                ai_style = st.selectbox(
                    "Report Style:",
                    ["Standard", "Detailed", "Concise", "Teaching"],
                    key="ai_style_select",
                    help="Choose the style of AI-generated report"
                )
                st.session_state.ai_style = ai_style
            
            elif ai_mode == "Enhance Existing":
                st.caption("AI will improve terminology and structure of your existing report")
            
            elif ai_mode == "Differential Diagnosis":
                st.caption("AI will provide differential diagnosis based on findings")
            
            elif ai_mode == "Generate Findings Only":
                st.caption("AI will expand brief descriptions into detailed findings")
                
            # Quick AI Input
            st.subheader("⚡ Quick AI Input")
            quick_prompt = st.text_area(
                "Brief description of findings:",
                height=100,
                placeholder="e.g., 'Right MCA territory infarct, mass effect, midline shift'",
                key="quick_ai_input"
            )
            
            col_ai1, col_ai2 = st.columns(2)
            with col_ai1:
                if st.button("🚀 Generate with AI", type="primary", use_container_width=True, 
                           disabled=not st.session_state.ai_helper.is_available()):
                    if quick_prompt:
                        with st.spinner("🤖 AI is generating report..."):
                            try:
                                if ai_mode == "Generate Full Report":
                                    ai_report = st.session_state.ai_helper.generate_report_from_findings(
                                        quick_prompt,
                                        modality=st.session_state.technique_info['modality'],
                                        contrast=st.session_state.technique_info['contrast'],
                                        style=st.session_state.ai_style
                                    )
                                    
                                elif ai_mode == "Generate Findings Only":
                                    ai_report = st.session_state.ai_helper.generate_findings_suggestions(quick_prompt)
                                    
                                elif ai_mode == "Differential Diagnosis":
                                    ai_report = st.session_state.ai_helper.generate_differential_diagnosis(quick_prompt)
                                    ai_report = f"DIFFERENTIAL DIAGNOSIS:\n{ai_report}"
                                    
                                else:  # Enhance Existing
                                    ai_report = st.session_state.ai_helper.enhance_existing_report(quick_prompt)
                                
                                if ai_report:
                                    # Clean and format AI response
                                    ai_report = clean_ai_response(ai_report)
                                    
                                    # Store in session state
                                    st.session_state.ai_report = ai_report
                                    st.session_state.report_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                                    st.session_state.ai_generated = True
                                    st.success("✅ AI report generated successfully!")
                                else:
                                    st.error("AI failed to generate report. Please try again.")
                                    
                            except Exception as e:
                                st.error(f"AI generation failed: {str(e)}")
                    else:
                        st.warning("Please enter some findings for AI to process")
            
            with col_ai2:
                if st.button("🔄 Test AI Connection", use_container_width=True,
                           help="Test if AI service is responding"):
                    if st.session_state.ai_helper.is_available():
                        with st.spinner("Testing connection..."):
                            try:
                                # Quick test query
                                test_response = st.session_state.ai_helper.client.query(
                                    model="sonar",
                                    messages=[
                                        {"role": "user", "content": "Say 'AI is working' if you can read this."}
                                    ],
                                    max_tokens=10
                                )
                                if test_response.choices:
                                    st.success("✅ AI connection successful!")
                                else:
                                    st.error("❌ AI connection test failed")
                            except Exception as e:
                                st.error(f"❌ Connection error: {str(e)}")
                    else:
                        st.error("❌ AI helper not initialized")
        else:
            # Show warning if AI is not available
            with st.container(border=True):
                st.warning("""
                **⚠️ AI Features Disabled**
                
                To enable AI features:
                1. Ensure you have a `.env` file with `PERPLEXITY_API_KEY`
                2. Your API key should be valid
                3. Restart the application
                
                File `.env` should contain:
                ```
                PERPLEXITY_API_KEY=your_api_key_here
                ```
                """)
        
        # Template Management
        st.divider()
        st.header("📚 Templates")
        
        # Add New Template
        with st.expander("➕ Add New Template", expanded=False):
            new_template_name = st.text_input(
                "Template Name", 
                placeholder="e.g., Normal Brain MRI Findings",
                key="new_template_name"
            )
            
            new_template_type = st.selectbox(
                "Template Type",
                ["findings", "technique", "impression"],
                format_func=lambda x: x.upper(),
                key="new_template_type"
            )
            
            new_template_content = st.text_area(
                "Template Content",
                height=150,
                placeholder="Enter the template text here...",
                key="new_template_content"
            )
            
            if st.button("💾 Save Template", type="primary", use_container_width=True, key="save_template"):
                if new_template_name and new_template_content:
                    st.session_state.template_system.add_template(
                        new_template_name,
                        new_template_content,
                        new_template_type
                    )
                    st.success(f"Template '{new_template_name}' saved!")
                    st.rerun()
                else:
                    st.warning("Please fill in both name and content")
        
        # Available Templates
        user_templates = st.session_state.template_system.get_user_templates(st.session_state.current_user)
        
        if user_templates:
            st.subheader("Your Templates:")
            for name, data in user_templates.items():
                col_temp1, col_temp2 = st.columns([3, 1])
                with col_temp1:
                    st.write(f"**{name}**")
                    st.caption(f"Type: {data.get('type', 'unknown').upper()} | Used: {data.get('used_count', 0)} times")
                with col_temp2:
                    if st.button("📥", key=f"insert_{name}", help=f"Insert {name}"):
                        st.session_state.report_draft = st.session_state.template_system.apply_template(
                            name,
                            st.session_state.report_draft
                        )
                        st.success(f"Template '{name}' inserted!")
                        st.rerun()
        else:
            st.info("No templates yet. Create your first template above!")
        
        # Quick Templates
        st.divider()
        st.subheader("⚡ Quick Templates")
        
        quick_templates = {
            "Normal Brain MRI": "No acute intracranial hemorrhage, mass effect, or territorial infarct. Ventricles and sulci are normal. No abnormal enhancement.",
            "White Matter Changes": "Scattered punctate FLAIR hyperintensities in the periventricular and deep white matter, consistent with chronic microvascular ischemic changes.",
            "Disc Herniation": "Disc bulge/protrusion causing mild neural foraminal narrowing without significant cord compression."
        }
        
        quick_selected = st.selectbox("Select:", ["Select..."] + list(quick_templates.keys()))
        
        if quick_selected != "Select...":
            if st.button(f"Insert '{quick_selected}'", use_container_width=True):
                st.session_state.report_draft += f"\nFINDINGS:\n{quick_templates[quick_selected]}"
                st.success("Template inserted!")
                st.rerun()
    
    with col2:
        # ===== RIGHT PANEL: DRAFTING & FINAL REPORT =====
        st.header("📝 Report Drafting")
        
        # Draft text area
        draft_text = st.text_area(
            "Type your report below:",
            value=st.session_state.report_draft,
            height=300,
            key="draft_input",
            label_visibility="collapsed",
            placeholder="Start typing your report here...\n\nUse 'FINDINGS:' and 'IMPRESSION:' headings.\n\nOr use AI assistance on the left."
        )
        st.session_state.report_draft = draft_text
        
        # Action buttons
        col_btn1, col_btn2, col_btn3, col_btn4 = st.columns(4)
        
        with col_btn1:
            if st.button("🤖 AI Generate", type="primary", use_container_width=True, 
                        help="Generate complete report using AI",
                        disabled=not (st.session_state.ai_helper and st.session_state.ai_helper.is_available())):
                if draft_text:
                    with st.spinner("🤖 AI is generating report..."):
                        try:
                            # Parse findings for AI
                            findings_for_ai = parse_findings_for_ai(draft_text)
                            
                            if not findings_for_ai:
                                findings_for_ai = draft_text
                            
                            # Generate AI report
                            ai_report = st.session_state.ai_helper.generate_report_from_findings(
                                findings_for_ai,
                                modality=st.session_state.technique_info['modality'],
                                contrast=st.session_state.technique_info['contrast'],
                                style=st.session_state.ai_style
                            )
                            
                            if ai_report:
                                # Clean and format
                                ai_report = clean_ai_response(ai_report)
                                
                                # Store in session state
                                st.session_state.ai_report = ai_report
                                st.session_state.report_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                                st.session_state.ai_generated = True
                                st.success("✅ AI report generated successfully!")
                            else:
                                st.error("AI failed to generate report")
                                
                        except Exception as e:
                            st.error(f"AI generation failed: {str(e)}")
                else:
                    st.warning("Please enter some findings first")
        
        with col_btn2:
            if st.button("✨ AI Enhance", use_container_width=True, 
                        help="Enhance existing report with AI",
                        disabled=not (st.session_state.ai_helper and st.session_state.ai_helper.is_available())):
                if draft_text:
                    with st.spinner("🤖 AI is enhancing report..."):
                        try:
                            enhanced = st.session_state.ai_helper.enhance_existing_report(draft_text)
                            if enhanced:
                                st.session_state.report_draft = enhanced
                                st.success("✅ Report enhanced with AI!")
                                st.rerun()
                            else:
                                st.error("AI enhancement failed")
                        except Exception as e:
                            st.error(f"AI enhancement failed: {str(e)}")
                else:
                    st.warning("Please enter a report to enhance")
        
        with col_btn3:
            if st.button("🔍 Check Structure", use_container_width=True):
                if "FINDINGS:" in draft_text.upper():
                    st.success("✓ FINDINGS section found")
                else:
                    st.warning("Missing FINDINGS section")
                
                if "IMPRESSION:" in draft_text.upper():
                    st.success("✓ IMPRESSION section found")
                else:
                    st.warning("Missing IMPRESSION section")
        
        with col_btn4:
            if st.button("🧹 Clear Draft", use_container_width=True):
                st.session_state.report_draft = ""
                st.rerun()
        
        # Final Report Display
        st.divider()
        st.header("📋 Final Report")
        
        if st.session_state.ai_report:
            # Display final report
            st.text_area(
                "Generated Report:",
                value=st.session_state.ai_report,
                height=350,
                key="report_preview",
                label_visibility="collapsed"
            )
            
            # Create Word document
            try:
                doc = create_word_document(
                    report_text=st.session_state.ai_report,
                    technique_info=st.session_state.technique_info,
                    report_date=st.session_state.report_date
                )
                
                # Save to buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Download button
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                
                st.download_button(
                    label="📄 Download as Word Document",
                    data=buffer,
                    file_name=f"RadReport_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )
                
            except Exception as e:
                st.error(f"Error creating document: {str(e)}")
            
            # Save to history
            st.divider()
            report_name = st.text_input(
                "Save as:",
                value=f"Report_{st.session_state.report_date.split()[0]}",
                key="report_name"
            )
            
            if st.button("💾 Save to History", use_container_width=True):
                history_entry = {
                    "name": report_name,
                    "date": st.session_state.report_date,
                    "technique_info": st.session_state.technique_info,
                    "report": st.session_state.ai_report,
                    "created_by": st.session_state.current_user,
                    "ai_generated": st.session_state.ai_generated
                }
                st.session_state.report_history.append(history_entry)
                st.success("Saved to history!")
        
        else:
            # Empty state
            with st.container(border=True):
                if st.session_state.ai_helper and st.session_state.ai_helper.is_available():
                    st.info("""
                    **Ready to generate your report!**
                    
                    **Options:**
                    1. **AI Generate** - Create complete report from findings
                    2. **AI Enhance** - Improve existing report
                    3. **Manual** - Type directly with templates
                    
                    **Required Structure:**
                    - TECHNIQUE section (auto-added)
                    - FINDINGS section (your observations)
                    - IMPRESSION section (conclusions)
                    
                    **💡 Tip:** Use the AI panel on the left for quick generation!
                    """)
                else:
                    st.info("""
                    **Ready to generate your report!**
                    
                    **Note:** AI features are currently disabled.
                    
                    **Manual Options:**
                    1. Use templates from the left panel
                    2. Type directly in the text area
                    3. Ensure proper structure with FINDINGS: and IMPRESSION:
                    
                    To enable AI, check your `.env` file and restart.
                    """)
    
    # ===== REPORT HISTORY =====
    st.divider()
    st.header("📜 Recent Reports")
    
    if st.session_state.report_history:
        for i, report in enumerate(reversed(st.session_state.report_history[-5:])):
            with st.expander(f"{report['name']} - {report['date']} {'🤖' if report.get('ai_generated') else ''}", expanded=False):
                col_h1, col_h2 = st.columns([1, 1])
                with col_h1:
                    if st.button(f"📥 Load", key=f"load_{i}", use_container_width=True):
                        st.session_state.technique_info = report['technique_info']
                        st.session_state.ai_report = report['report']
                        st.session_state.report_date = report['date']
                        st.session_state.ai_generated = report.get('ai_generated', False)
                        st.success("Report loaded!")
                        st.rerun()
                
                with col_h2:
                    report_idx = len(st.session_state.report_history) - 1 - i
                    if st.button(f"🗑️ Delete", key=f"delete_{i}", use_container_width=True):
                        st.session_state.report_history.pop(report_idx)
                        st.warning("Report deleted!")
                        st.rerun()
                
                # Show preview
                preview = report['report'][:200] + "..." if len(report['report']) > 200 else report['report']
                st.text(preview)
    else:
        st.info("No reports in history yet. Generate and save your first report!")

# ===== RUN APPLICATION =====
if __name__ == "__main__":
    main()
